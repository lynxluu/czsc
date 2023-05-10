import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from threading import Lock, Thread
import io
from datetime import datetime


# 定义爬取的网站地址和页面范围
base_url = 'https://chzhshch.blog/stocks/'
page_range = range(1, 6)

# 创建一个列表，用于保存爬取的文章内容
articles = []

# 创建一个线程锁
lock = Lock()

def get_abs_url(rel_url):
    base_url = 'https://chzhshch.blog'
    url = base_url + rel_url
    return url

# 定义一个函数，用于爬取文章内容
def crawl(page_number):
    page_url = base_url + str(page_number).zfill(3)
    try:
        response = requests.get(page_url)
        response.raise_for_status()  # 检查响应状态码是否为 200
    except requests.exceptions.RequestException as e:
        print(f"请求页面 {page_url} 失败: {e}")
        return

    soup = BeautifulSoup(response.content, 'html.parser')
    article = soup.find('article')
    title = article.find('h1').text.strip()

    # 将爬取的文章内容保存到列表中
    with lock:
        articles.append((page_number,title, article.contents))

# 创建线程池，并启动多个线程进行爬取操作
threads = []
for page_number in page_range:
    thread = Thread(target=crawl, args=(page_number,))
    threads.append(thread)
    thread.start()

# 等待所有线程结束，然后对文章内容按照标题进行排序
for thread in threads:
    thread.join()

articles = sorted(articles, key=lambda x: x[0])

# 创建一个Word文档对象
document = Document()

# 将爬取的文章内容写入Word文档
for page_number, title, contents in articles:
    document.add_heading(title, level=1)
    # print(type(contents))
    for element in contents[0]:
        # print(element)
        if element.name == 'p':
            document.add_paragraph(element.text.strip())
        elif element.name == 'div':
            for img in element.children:
                if img.name == 'img':
                    image_url = img['src']
                    try:
                        # image_response = requests.get(image_url)
                        image_response = requests.get(get_abs_url(image_url),stream=True)
                        image = io.BytesIO(image_response.content)
                        image_response.raise_for_status()  # 检查响应状态码是否为 200
                    except requests.exceptions.RequestException as e:
                        print(f"请求图片 {image_url} 失败: {e}")
                    document.add_picture(image, width=Inches(6))

        # https://chzhshch.blog/stocks/002
        # 从这篇来看 h2后面接的平级的若干个div才是真正的回复
        elif element.name == 'h2' or element.name == 'h3':
            print(element)
            document.add_heading('回复', level=2)
            # 获取下一级标签是 <h2> 的内容
            for div in element.children:
                if div.name == 'div':
                    print('div',div)
                    # 遍历所有的 <div> 标签
                    # 获取回复者
                    replyer = ''
                    replyer_tag = div.find('span')
                    replyer = replyer_tag.get_text() if replyer_tag else ''
                    print(replyer)

                    # 获取回复时间
                    time_str = ''
                    try:
                        time_str = div.find('br').previous_sibling.strip()
                        reply_time = datetime.strptime(time_str, '%Y-%m-%d %H:%M:%S')
                    except Exception as e:
                        reply_time= None
                        print(e)
                    print(time_str)

                    document.add_paragraph(replyer+time_str)

                    #获取回复内容
                    reply_text = ''
                    try:
                        p_tags = div.find_all('p')
                        print(div,p_tags)
                        for p_tag in p_tags:
                            reply_text += p_tag.text.strip() + '\n'
                    except Exception as e:
                        print(e)
                    print(reply_text)

                    document.add_paragraph(reply_text)


            # for reply in element.find_all('span'):
            #     reply_text = ''
            #     for reply_element in reply.contents:
            #         if reply_element.name == 'p':
            #             reply_text += reply_element.text.strip() + '\n'
            #         elif reply_element.name == 'br':
            #             reply_text += '\n'
            #     document.add_paragraph(reply_text)

    # if contents[0].find('h2'):
    #     document.add_heading('回复', level=2)
    #     for reply in contents[0].find('h2').find_all('span'):
    #         reply_text = ''
    #         for reply_element in reply.contents:
    #             if reply_element.name == 'p':
    #                 reply_text += reply_element.text.strip() + '\n'
    #             elif reply_element.name == 'br':
    #                 reply_text += '\n'
    #         document.add_paragraph(reply_text)

    # print(page_number, title, contents)

# 保存Word文档
document.save('articles.docx')