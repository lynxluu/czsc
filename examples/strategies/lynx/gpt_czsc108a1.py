import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

# 发送请求
url = 'https://chzhshch.blog/stocks/005'
response = requests.get(url)

# 解析HTML
soup = BeautifulSoup(response.text, 'html.parser')
article = soup.find('article')

# 创建Word文档
document = Document()

# 遍历所有元素
for element in article.children:
    # 如果是p标签，就将其文本写入Word文档
    if element.name == 'p':
        document.add_paragraph(element.get_text())
    # 如果是img标签，就将其图片写入Word文档
    elif element.name == 'img':
        image_url = element['src']
        image_response = requests.get(image_url)
        image_file = BytesIO(image_response.content)
        document.add_picture(image_file, width=Inches(6))

# 保存Word文档
document.save('output.docx')


按顺序爬取https://chzhshch.blog/stocks/005这个页面的article元素里的p标签，img标签等所有元素并写入到word文档

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

# 发送请求
url = 'https://chzhshch.blog/stocks/005'
response = requests.get(url)

# 解析HTML
soup = BeautifulSoup(response.text, 'html.parser')
article = soup.find('article')

# 创建Word文档
document = Document()

# 遍历所有元素
for element in article.children:
    # 如果是p标签，就将其文本写入Word文档
    if element.name == 'p':
        document.add_paragraph(element.get_text())
    # 如果是img标签，就将其图片写入Word文档
    elif element.name == 'img':
        image_url = element['src']
        image_response = requests.get(image_url)
        image_file = BytesIO(image_response.content)
        document.add_picture(image_file, width=Inches(6))

# 保存Word文档
document.save('output.docx')
