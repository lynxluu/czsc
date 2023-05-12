import os
# import re
import io
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from docx import Document
# from zipfile import ZipFile
# from io import BytesIO

from docx.shared import Pt,Cm,Inches
from docx.oxml.ns import qn
# from docx.enum.section import WD_SECTION
from docx.image.image import Image


import concurrent.futures

import urllib3
# from lxml import etree

# 忽略 SSL 证书警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)



def get_content2c(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # 检查响应状态码是否为 200
    except requests.exceptions.RequestException as e:
        print(f"请求页面 {url} 失败: {e}")
        return

    soup = BeautifulSoup(response.content, 'html.parser')
    article = soup.find('article')
    # 获取标题
    title_tag = article.find('h1')
    title = title_tag.get_text() if title_tag else ''

    # 获取时间
    time_tag = article.find('blockquote')
    time_str = time_tag.get_text() if time_tag else ''
    time = datetime.strptime(time_str, '%Y/%m/%d %H:%M:%S')

    # 获取内容
    content = article.contents

    document = Document()

    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 添加标题
    document.add_heading(title, level=1)
    # print(type(contents))

    # 添加段落，图片
    # print(content)
    for element in content[0]:
        print("element",element)
        if element.name == 'p':
            document.add_paragraph(element.text.strip())
        elif element.name == 'div':
            for img in element.children:
                if img.name == 'img':
                    image_url = img['src']
                    try:
                        # image_response = requests.get(image_url)
                        image_response = requests.get(get_abs_url(image_url), stream=True)
                        image = io.BytesIO(image_response.content)
                        document.add_picture(image, width=Inches(6))
                        # print(get_abs_url(image_url))
                    except requests.exceptions.RequestException as e:
                        print(f"请求图片 {image_url} 失败: {e}")

        # 添加回复
        elif element.name in ['h2', 'h3']:
            # print('h2',element)
            document.add_heading('回复', level=2)
            # 获取下一级标签是 <h2> 的内容
            divs = element.find_next_siblings('div')
            for div in divs:
                # 获取回复者
                replyer = ''
                replyer_tag = div.find('span')
                replyer = replyer_tag.get_text() if replyer_tag else ''

                # 获取回复时间
                time_str = ''
                try:
                    time_str = div.find('br').previous_sibling.strip()
                    # reply_time = datetime.strptime(time_str, '%Y-%m-%d %H:%M:%S')
                    if not time_str:
                        time_str = None
                except Exception as e:
                    print(e)
                # print(time_str)
                document.add_paragraph(replyer+': ' +time_str)

                #获取回复内容
                reply_text = ''
                try:
                    p_tags = div.find_all('p')
                    # 如果p_tags存在，循环拼接p标签内容；如不存在，则取div内容
                    if p_tags:
                        # print(div,p_tags)
                        for p_tag in p_tags:
                            reply_text += p_tag.text.strip() + '\n'
                    else:
                        #替换掉多余的回复人和、回复时间
                        reply_text = div.get_text(strip=True).replace(time_str, '').replace(replyer, '') + '\n'
                except Exception as e:
                    print(e)

                # print(reply_text)
                document.add_paragraph(reply_text)

    # 添加分页符，防止格式丢失
    # document.add_page_break()
    # document.add_section()

    return [title, time, article,document]



def save_to_word2c(urls, file_path):
    # 定义一个空的内容列表
    contents = []

    def process_url(url):
        # 获取内容
        print(f"processing {url}")
        # 将文档对象保存到内容列表中
        contents.append(get_content2c(url))

    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        futures = [executor.submit(process_url, url) for url in urls]
        for future in concurrent.futures.as_completed(futures):
            try:
                future.result()
            except Exception as e:
                print(f"Error processing URL: {e}")

    # 根据时间对内容列表进行排序
    # for each in contents: print(each[0], each[1])
    contents.sort(key=lambda x: x[1])
    for each in contents: print(each[0], each[1])

    # 将排序后的内容写入 Word 文件
    merged_document = Document()

    # if os.path.exists(file_path):
    #     merged_document = Document(file_path)
    # else:
    #     merged_document = Document()

    section = merged_document.sections[0]
    section.page_width = Inches(6)
    # merged_document.default_width = Inches(6)
    merged_document.styles['Normal'].font.name = '微软雅黑'
    merged_document.styles['Normal'].font.size = Pt(12)
    merged_document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for _, _, _, content_doc in contents:
        for element in content_doc.element.body:
            if isinstance(element, Image):
                print(element.attrib)
            merged_document.element.body.append(element)
            # merged_document.add_document(content_doc)
        # merged_document.add_page_break()

    merged_document.save(file_path)
    # merge_documents(contents,file_path)



def get_abs_url(rel_url):
    base_url = 'https://chzhshch.blog'
    url = base_url + rel_url
    return url


def get_links(url):
    # 定义为集合，避免重复
    links = set()
    # print(url)
    response = requests.get(url, verify=False, proxies=None)
    if response.status_code != 200:
        return links

    soup = BeautifulSoup(response.content, 'html.parser')

    # 寻找链接的标签，这里需要根据目标网站的实际页面结构进行修改
    anchor_tags = soup.find_all('a',href=True)

    for tag in anchor_tags:
        link = tag['href']
        if link.startswith('/stocks/'):  # 只获取以 '/stocks/' 开头的链接
            links.add(link)

    # lists排序
    sorted_links = sorted(links)
    # return sorted_links[:109]
    return sorted_links

# 5线程
def main():
    target_url = get_abs_url('/stocks/wolves')
    # file_path = os.path.join(r"D:\usr\doc", 'czsc108' + '.docx')
    file_path = os.path.join('czsc108text.docx')
    # print(target_url)
    # 获取所有相关相对链接
    # rel_urls = get_links(target_url)[:110]
    rel_urls = get_links(target_url)[6:7]
    urls = []
    # print(rel_urls)

    # 获得绝对链接
    for rel_url in rel_urls:
        url = get_abs_url(rel_url)
        # save_to_word3(url, file_path)
        urls.append(url)

    save_to_word2c(urls, file_path)

if __name__ == '__main__':
    main()