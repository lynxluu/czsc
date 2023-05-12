import os
# import re
import io
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from docx import Document
# from zipfile import ZipFile
# from io import BytesIO

from docx.shared import Pt,Cm,Inches
from docx.oxml.ns import qn
# from docx.enum.section import WD_SECTION
from docx.image.image import Image


import concurrent.futures

import urllib3
# from lxml import etree

# 忽略 SSL 证书警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


def get_abs_url(rel_url):
    base_url = 'https://chzhshch.blog'
    url = base_url + rel_url
    return url


def get_links(url):
    # 定义为集合，避免重复
    links = set()
    # print(url)
    response = requests.get(url, verify=False, proxies=None)
    if response.status_code != 200:
        return links

    soup = BeautifulSoup(response.content, 'html.parser')

    # 寻找链接的标签，这里需要根据目标网站的实际页面结构进行修改
    anchor_tags = soup.find_all('a',href=True)

    for tag in anchor_tags:
        link = tag['href']
        if link.startswith('/stocks/'):  # 只获取以 '/stocks/' 开头的链接
            links.add(link)

    # lists排序
    sorted_links = sorted(links)
    # return sorted_links[:109]
    # return sorted_links[:5]
    return sorted_links

# 单线程循环
def main():
    target_url = get_abs_url('/stocks/wolves')
    # file_path = os.path.join(r"D:\usr\doc", 'czsc108' + '.docx')
    file_path = os.path.join('czsc108.docx')
    # print(target_url)
    # 获取所有相关相对链接
    # rel_urls = get_links(target_url)[:110]
    rel_urls = get_links(target_url)[79:85]
    urls = []
    # print(rel_urls)

    document = Document()

    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 获得绝对链接
    for rel_url in rel_urls:
        url = get_abs_url(rel_url)
        # save_to_word3(url, file_path)
        urls.append(url)

        try:
            print(f"processing {url}")
            response = requests.get(url)
            response.raise_for_status()  # 检查响应状态码是否为 200
        except requests.exceptions.RequestException as e:
            print(f"请求页面 {url} 失败: {e}")
            return

        soup = BeautifulSoup(response.content, 'html.parser')
        article = soup.find('article')
        content = article.contents

        # 获取标题
        title_tag = article.find('h1')
        title = title_tag.get_text() if title_tag else ''
        document.add_heading(title, level=1)

        # 获取时间
        time_tag = article.find('blockquote')
        time_str = time_tag.get_text() if time_tag else ''
        time = datetime.strptime(time_str, '%Y/%m/%d %H:%M:%S')

        # 递归遍历所有子标签
        def traverse(tag):
            if tag.name is None:
                return
            for child in tag.children:
                if child.name == 'p':
                    document.add_paragraph(child.text)
                elif child.name == 'img':
                    image_url = child['src']
                    try:
                        image_response = requests.get(get_abs_url(image_url), stream=True)
                        image = io.BytesIO(image_response.content)
                        document.add_picture(image, width=Inches(6))
                    except requests.exceptions.RequestException as e:
                        print(f"请求图片 {image_url} 失败: {e}")
                elif child.name in ['h2', 'h3']:
                    document.add_heading('回复', level=2)
                elif child.name == 'span':
                    document.add_paragraph(child.text)
                elif child.name =='br':
                    document.add_paragraph(child.previous_sibling.text)
                else:
                    traverse(child)

        # 遍历正文
        for tag in article:
            traverse(tag)

    document.save(file_path)
    # save_to_word2c(urls, file_path)

if __name__ == '__main__':
    main()