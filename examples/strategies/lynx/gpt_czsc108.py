import os

import requests
from bs4 import BeautifulSoup
from docx import Document
import urllib3

# 忽略 SSL 证书警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# 获取网页链接和所需的数据
def get_links(url):
    # 定义为集合，避免重复
    links = set()
    # print(url)
    response = requests.get(url, verify=False, proxies=None)
    if response.status_code != 200:
        return links

    soup = BeautifulSoup(response.content, 'html.parser')

    # 寻找链接的标签，这里需要根据目标网站的实际页面结构进行修改
    anchor_tags = soup.find_all('a',href=True)

    for tag in anchor_tags:
        link = tag['href']
        if link.startswith('/stocks/'):  # 只获取以 '/stocks/' 开头的链接
            links.add(link)

    # lists排序
    sorted_links = sorted(links)
    return sorted_links[:5]


# 爬取特定链接内容
def get_content(url):
    response = requests.get(url, verify=False, proxies=None)
    res = []
    if response.status_code != 200:
        return res
    soup = BeautifulSoup(response.content, 'html.parser')

    # 找到页面中包含文章内容的部分，需要根据实际页面结构进行修改
    title_tag = soup.find('h1')
    if title_tag:
        title = title_tag.get_text()

    time_tag = soup.find('blockquote')
    if time_tag:
        time = time_tag.get_text()

    content_tag = soup.find('article')
    if content_tag:
        content = content_tag.get_text()

    res = [title,time,content]
    return res


# 保存内容到单个 Word 文档
def save_to_word(doc, content, title):
    doc.add_heading(title, level=1)
    # 添加正文内容到 Word 文档
    doc.add_paragraph(content)

def main():
    base_url = 'https://chzhshch.blog'
    target_url = base_url + '/stocks/wolves'
    # print(target_url)
    # 获取所有相关链接
    links = get_links(target_url)

    # 创建一个新的 Word 文档，用于保存所有文章内容
    doc = Document()

    # 遍历链接，获取文章内容并保存到 Word 文档
    for link in links:
        url = base_url + link
        res = get_content(url)
        # print(url, res)

        title = res[0]
        time = res[1]
        content = res[2]

        print(f"Processing {link,title}")
        save_to_word(doc, content, title)


    # 指定保存路径
    file_path = os.path.join(r"D:\usr\doc",'czsc108' + '.docx')
    doc.save(file_path)

if __name__ == '__main__':
    main()

# base_url = 'https://chzhshch.blog'
# target_url = base_url + '/stocks/wolves'
# print(target_url)
# # 获取所有相关链接
# links = get_links(target_url)
# print(links)