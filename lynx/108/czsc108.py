import os
import re
import io

import requests
from bs4 import BeautifulSoup
from docx import Document
from zipfile import ZipFile
from io import BytesIO

from docx.shared import Pt
from docx.oxml.ns import qn
from PIL import Image
from docx.shared import Cm
import urllib3
from lxml import etree

# 忽略 SSL 证书警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# 获取网页链接和所需的数据
def get_links(url):
    # 定义为集合，避免重复
    links = set()
    # print(url)
    response = requests.get(url, verify=False, proxies=None)
    if response.status_code != 200:
        return links

    soup = BeautifulSoup(response.content, 'html.parser')

    # 寻找链接的标签，这里需要根据目标网站的实际页面结构进行修改
    anchor_tags = soup.find_all('a',href=True)

    for tag in anchor_tags:
        link = tag['href']
        if link.startswith('/stocks/'):  # 只获取以 '/stocks/' 开头的链接
            links.add(link)

    # lists排序
    sorted_links = sorted(links)
    return sorted_links[:109]


# 爬取特定链接内容
def get_content(url):
    response = requests.get(url, verify=False, proxies=None)
    res = []
    if response.status_code != 200:
        return res
    soup = BeautifulSoup(response.content, 'html.parser')

    # 找到页面中包含文章内容的部分，需要根据实际页面结构进行修改
    title_tag = soup.find('h1')
    if title_tag:
        title = title_tag.get_text()

    time_tag = soup.find('blockquote')
    if time_tag:
        time = time_tag.get_text()

    content_tag = soup.find('article')
    if content_tag:
        content = content_tag.get_text()

    res = [title,time,content]
    return res

def get_content2(url):
    try:
        # 发送HTTP请求
        response = requests.get(url)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(e)
        return []

    # 解析HTML文档
    soup = BeautifulSoup(response.content, 'html.parser')

    # 获取标题
    title_tag = soup.find('h1')
    title = title_tag.get_text() if title_tag else ''

    # 获取时间
    time_tag = soup.find('blockquote')
    time = time_tag.get_text() if time_tag else ''

    # 获取内容
    content_tag = soup.find('article')
    content = ''
    if content_tag:
        # 获取段落
        paragraphs = content_tag.find_all('p')
        content = '\n'.join([p.get_text() for p in paragraphs])

        # 获取图片
        images = content_tag.find_all('img')
        for img in images:
            src = img.get('src')
            alt = img.get('alt')
            size = (img.get('width'), img.get('height'))
            content += f'\n[Image] src={src} alt={alt} size={size}'

    return [title, time, content]

# 保存内容到单个 Word 文档
def save_to_word(doc, content, title):
    doc.add_heading(title, level=1)
    # 添加正文内容到 Word 文档
    doc.add_paragraph(content)

# 保存多个url地址到同一个word文档
def save_to_word2(urls, file_path):
    # 创建Word文档
    document = Document()

    # 设置字体
    document.styles['Normal'].font.name = 'Microsoft YaHei'
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    # 将页面边距设置为A4纸的大小
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    # 设置图片宽度
    width = section.page_width - section.left_margin - section.right_margin


    # 循环处理每个网页
    for url in urls:
        print(f"processing {url}")
        # 获取内容
        content = get_content2(url)
        # print(content)

        # 添加标题
        title = content[0]
        document.add_heading(title, level=1)

        # # 添加时间
        # time = content[1]
        # document.add_paragraph(time)

        # 添加内容
        sub_content = content[2]
        # document.add_paragraph(sub_content)
        paragraphs = sub_content.split('\n')
        for p in paragraphs:
            if p.startswith('[Image]'):
                # 添加图片
                # print(p,p.split())
                _, src, alt, size = p.split(' ', 3)
                src = src[4:]
                response = requests.get(get_abs_url(src), stream=True)
                pic = response.content
                file_obj = io.BytesIO(pic)

                # 缩放图片并将其添加到Word文档中,指定宽度，自动缩放，需要设置为A4纸的宽度大约为16cm
                # 原来设置width=500超出了页面限制，所以无法显示。
                # document.add_picture(file_obj)
                document.add_picture(file_obj, width=width)

            else:
                # 添加段落
                paragraph = document.add_paragraph()
                paragraph.add_run(p).font.size = Pt(12)

    # 保存文档
    document.save(file_path)

def save_to_word3(url, file_path):
    # 不存在则创建Word文档
    if not os.path.isfile(file_path):
        document = Document()
        # 添加一个标题
        document.add_heading('Document Title', level=0)
        # 添加一个段落
        document.add_paragraph('This is a paragraph.')
        document.save(file_path)

    # 否则打开现有的Word文档
    with ZipFile(file_path, 'a') as zip:
        # 读取Word文档中的document.xml文件
        with zip.open('word/document.xml') as xml:
            # 将文件内容读取到内存中
            xml_content = xml.read()

        # 将文件内容加载到Document对象中
        # 将XML内容解析为元素树
        root = etree.fromstring(xml_content)

        # 创建一个新的Document对象
        document = Document()

        # 将元素树加载到Document对象中
        document_part = document.part
        document_part._element.clear()
        document_part._element.append(root)

        # 设置字体
        document.styles['Normal'].font.name = 'Microsoft YaHei'
        document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 处理网页
        print(f"processing {url}")
        # 获取内容
        content = get_content2(url)
        # print(content)

        # 添加标题
        title = content[0]
        document.add_heading(title, level=1)


        # 添加内容
        sub_content = content[2]
        # document.add_paragraph(sub_content)
        paragraphs = sub_content.split('\n')
        for p in paragraphs:
            if p.startswith('[Image]'):
                # 添加图片
                # print(p,p.split())
                _, src, alt, size = p.split(' ', 3)
                src = src[4:]
                response = requests.get(get_abs_url(src), stream=True)
                pic = response.content
                file_obj = io.BytesIO(pic)

                document.add_picture(file_obj, width=16)

            else:
                # 添加段落
                paragraph = document.add_paragraph()
                paragraph.add_run(p).font.size = Pt(12)

        # 将文档保存回原始文件
        with zip.open('word/document.xml', 'w') as xml:
            xml.write(document._element.xml.encode())

    # # 保存文档
    # document.save(filename)
def get_abs_url(rel_url):
    base_url = 'https://chzhshch.blog'
    url = base_url + rel_url
    return url

def main():
    target_url = get_abs_url('/stocks/wolves')
    file_path = os.path.join(r"D:\usr\doc",'czsc108' + '.docx')
    # print(target_url)
    # 获取所有相关相对链接
    rel_urls = get_links(target_url)
    urls = []
    # print(rel_urls)

    # 获得绝对链接
    for rel_url in rel_urls:
        url = get_abs_url(rel_url)
        save_to_word3(url, file_path)
        # urls.append(url)
        # urls.append(get_abs_url(rel_url))

    # 指定保存路径
    # save_to_word2(urls, file_path)

if __name__ == '__main__':
    main()

# base_url = 'https://chzhshch.blog'
# target_url = base_url + '/stocks/wolves'
# print(target_url)
# # 获取所有相关链接
# links = get_links(target_url)
# print(links)