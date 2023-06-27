from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_toc(document, headings):
    # 创建一个空的表格，用于存放目录
    toc = document.add_table(rows=1, cols=2)
    toc.style = 'Table Grid'

    # 设置目录标题和字体样式
    heading = toc.cell(0, 0).paragraphs[0]
    heading.text = '目录'
    heading.style = document.styles['Heading 1']
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run()
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(24)

    # 添加目录内容
    for title, page in headings.items():
        paragraph = toc.add_row().cells[0].paragraphs[0]
        # hyperlink = paragraph.add_hyperlink('')
        # hyperlink = paragraph.add_hyperlink('')
        # hyperlink.text = title
        # hyperlink.style = document.styles['Hyperlink']
        # hyperlink.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # run = hyperlink.add_run(f'\t{page}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # 更新目录
    document.add_page_break()
    # document.update_toc()

# 给定的字典
headings = {'第一章': '005', '第二章': '020', '第三章': '050'}

# 创建一个新文档，并添加章节标题和内容
document = Document()
document.add_heading('第一章', level=1)
document.add_paragraph('第一章内容')
document.add_heading('第二章', level=1)
document.add_paragraph('第二章内容')
document.add_heading('第三章', level=1)
document.add_paragraph('第三章内容')

# 添加目录
add_toc(document, headings)

# 保存文档
document.save('test.docx')