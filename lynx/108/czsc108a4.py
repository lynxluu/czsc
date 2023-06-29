import os
import threading
# import io
from datetime import datetime
from lib108 import *

import requests
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Pt,Cm,Inches
from docx.oxml.ns import qn
from docx.image.image import Image


# import concurrent.futures


# 忽略 SSL 证书警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


# 单线程循环, 解析html使用递归，直接写文件
def single_doc():
    target_url = get_abs_url('/stocks/wolves')

    ext = datetime.now().strftime('%y%m%d%H%M%S')
    file_path = os.path.join(f'czsc108a3_{ext}.docx')
    print(file_path)


    # 获取所有相关相对链接
    rel_urls = get_links(target_url)[7:9]

    document = Document()
    set_format(document)

    urls = []
    # 获得绝对链接
    for rel_url in rel_urls:
        url = get_abs_url(rel_url)
        urls.append(url)

        # 获取对应url的内容
        response = get_response(url)

        soup = BeautifulSoup(response.content, 'html.parser')
        article = soup.find('article')
        # content = article.contents

        # 获取标题
        title_tag = article.find('h1')
        title = title_tag.get_text() if title_tag else ''
        document.add_heading(title, level=1)

        # 获取时间
        # time_tag = article.find('blockquote')
        # time_str = time_tag.get_text() if time_tag else ''
        # time = datetime.strptime(time_str, '%Y/%m/%d %H:%M:%S')
        # document.add_paragraph(time_str)

        # 递归遍历所有子标签
        def traverse(tag):
            if tag.name is None:
                return
            for child in tag.children:
                if child.name == 'p':
                    document.add_paragraph(child.text)
                elif child.name == 'img':
                    image_url = get_abs_url(child['src'])
                    image = get_image(image_url)
                    if image:
                        document.add_picture(image, width=Inches(6))

                elif child.name in ['h2', 'h3']:
                    document.add_paragraph('回复')
                elif child.name == 'span':
                    document.add_paragraph(child.text)
                elif child.name =='br':
                    # 杜绝无内容的br占据很大篇幅
                    if child.previous_sibling.name != 'br':
                        document.add_paragraph(child.previous_sibling.text)
                else:
                    traverse(child)

        # 遍历正文
        for tag in article:
            traverse(tag)

    document.save(file_path)


def get_range(rel_urls, img_flag=False):
    f_start = rel_urls[0].split("/")[-1]
    f_end = rel_urls[-1].split("/")[-1]
    ext = datetime.now().strftime('%y%m%d%H%M%S')
    file_path = os.path.join(f'czsc108a4_{f_start}_{f_end}_{ext}.docx')

    document = Document()
    set_format(document)

    urls = []
    # 获得绝对链接
    for rel_url in rel_urls:
        url = get_abs_url(rel_url)
        urls.append(url)

        response = get_response(url)

        if not response:
            return

        soup = BeautifulSoup(response.content, 'html.parser')
        article = soup.find('article')
        # content = article.contents

        # 获取标题
        title_tag = article.find('h1')
        title = title_tag.get_text() if title_tag else ''
        document.add_heading(title, level=1)

        # 获取时间
        # time_tag = article.find('blockquote')
        # time_str = time_tag.get_text() if time_tag else ''
        # time = datetime.strptime(time_str, '%Y/%m/%d %H:%M:%S')
        # document.add_paragraph(time_str)

        # 递归遍历所有子标签
        def traverse(tag):
            if tag.name is None:
                return
            for child in tag.children:
                if child.name == 'p':
                    document.add_paragraph(child.text)
                elif child.name == 'img':
                    image_url = get_abs_url(child['src'])
                    if not img_flag:
                        document.add_paragraph(image_url)
                    else:
                        add_image(image_url, document)

                elif child.name in ['h2', 'h3']:
                    # document.add_heading('回复', level=2)
                    document.add_paragraph('回复')
                elif child.name == 'span':
                    document.add_paragraph(child.text)
                elif child.name == 'br':
                    # 杜绝无内容的br占据很大篇幅
                    if child.previous_sibling.name != 'br':
                        document.add_paragraph(child.previous_sibling.text)
                else:
                    traverse(child)

        # 遍历正文
        for tag in article:
            traverse(tag)

    document.save(file_path)


def list_doc():
    target_url = get_abs_url('/stocks/wolves')

    ext = datetime.now().strftime('%y%m%d%H%M%S')
    file_path = os.path.join(f'czsc108a4_{ext}.docx')
    print(file_path)

    # 获取所有相关相对链接
    rel_urls = get_links(target_url)[7:9]

    document = Document()
    set_format(document)

    urls = []
    pages = []
    # 获得绝对链接
    for rel_url in rel_urls:
        url = get_abs_url(rel_url)
        urls.append(url)

        # 获取对应url的内容
        response = get_response(url)

        soup = BeautifulSoup(response.content, 'html.parser')
        article = soup.find('article')
        # content = article.contents

        # 查找文章标题和发表时间
        title = article.find('h1').text.strip()
        ptime = article.find('blockquote').text.strip()

# 定义线程的执行函数
def multi_run(rel_url_list):
    def process_rel_url(sublist):
        get_range(sublist)

    # 创建5个线程并启动
    threads = []
    for sublist in rel_url_list:
        t = threading.Thread(target=process_rel_url, args=(sublist,))
        threads.append(t)
        t.start()

    # 等待所有线程执行完成
    for t in threads:
        t.join()
def main():
    target_url = get_abs_url('/stocks/wolves')
    rel_urls = get_links(target_url)[:109]

    n = 20  # 切片大小为20
    rel_url_list = [rel_urls[i:i + n] for i in range(0, len(rel_urls), n)]
    # print(res)

    #顺序执行
    # for each in rel_url_list:
    #     get_range(each)

    #执行单个
    get_range(rel_urls[60:80])
    # get_range(rel_urls[66:67])

    # 多线程执行
    # multi_run(rel_url_list)


if __name__ == '__main__':
    main()



