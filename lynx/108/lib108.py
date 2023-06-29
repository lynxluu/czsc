import requests
import urllib3
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO
from docx.shared import Pt,Cm,Inches
from docx.oxml.ns import qn
from PIL import Image
import math

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

"""
已实现的函数说明
czsc108c1：多线程ok，图片不能显示
czsc108a3：多线程ok，图片正常，文章日期有两个
"""

"""
gpt提示词
现在要爬取一个网站，网站的每个页面地址都是https://chzhshch.blog/stocks/XXX，
其中XXX代表001-108，每个页面的文章是用article标签框定，
article中有h1标签代表文章的标题，blockquote标签代表发表时间，
接着就是p标签代表一段文字，img标签代表这段文字对应的图片，有时候有文字有图片，有时候仅有问题没有图片。
然后文章快结束时，有可能出现h2或h3,是代表回复的内容。一个div标签代表一条回复，多个div代表多个回复。
每个回复的内容中，span标签代表回复人，第一个<br/>标签前面代表回复时间，其余的br标签和p标签就是回复内容。
将以上内容爬取后，生成一份word,并按照各篇文章的标题形成目录。

# 会以span标签进行回复项的分割，在每个span中会以p标签或者br标签框定文字，

# 1） 注意，就算请求时使用了 proxies=None 参数，还是要关闭ie的默认代理才能访问
response = requests.get(url, headers=headers, verify=False, proxies=None)
Max retries exceeded with url: /stocks/001 (Caused by ProxyError('Cannot connect to proxy.', 
FileNotFoundError(2, 'No such file or directory')))

# 2）需要忽略 SSL 证书警告
InsecureRequestWarning: Unverified HTTPS request is being made to host 'chzhshch.blog'. 
Adding certificate verification is strongly advised

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# 3) DPI计算
PPI = sqrt(水平像素数^2 + 垂直像素数^2) / 对角线长度（英寸）
# 获取图片信息
image_size = len(image.getbuffer())
width, height = im.size
# print(width,height,image_size,dpi)
# print(im.size,len(image.getbuffer()),dpi)
"""

# 根据相对地址获取可以访问的绝对地址
def get_abs_url(rel_url):
    # print(f"------get_abs_url:{rel_url}")
    base_url = 'https://chzhshch.blog'
    url = base_url + rel_url
    return url

# 获取根站点下所有网页
def get_links(url):
    print(f"------get_links:{url}")
    # 定义为集合，避免重复
    links = set()
    response = None
    try:
        response = requests.get(url, verify=False, proxies=None)
    except Exception as e:
        print(f"http请求错误：{e}")

    if not response or response.status_code != 200:
        return []

    soup = BeautifulSoup(response.content, 'html.parser')

    # 寻找链接的标签，这里需要根据目标网站的实际页面结构进行修改
    anchor_tags = soup.find_all('a',href=True)

    for tag in anchor_tags:
        link = tag['href']
        if link.startswith('/stocks/'):  # 只获取以 '/stocks/' 开头的链接
            links.add(link)

    # lists排序
    sorted_links = sorted(links)
    return sorted_links


def add_image(url, document):
    image = get_image(url)
    if image:
        if get_image_dpi(image) > 0:
            document.add_picture(image, width=Inches(6))
        else:
            # print(type(set_image_dpi(image)), type(image))
            print(f"图片没有dpi信息，设默认dpi为96:{url}")
            document.add_picture(set_image_dpi(image), width=Inches(6))
    return


def get_image_dpi(image):
    # 判断图像的尺寸
    with Image.open(image) as im:
        dpi = im.info.get('dpi', (0, 0))[0]
        if dpi == 0:
            return 96
        return dpi


def set_image_dpi(image):
    # 判断图像的尺寸
    with Image.open(image) as im:
        # 获取图片文件的后缀名
        suffix = im.format.lower()

        dpi = im.info.get('dpi', (0, 0))
        if dpi[0] == 0 or dpi[1] == 0:
            # DPI无法计算，给个默认值吧
            im.info['dpi'] =(96,96)

        # print(im.size, len(image.getbuffer()), dpi)

        # 将 DPI 修改后的图片对象转换为二进制数据,存储到 BytesIO 对象,并返回
        buffered = BytesIO()
        im.save(buffered, format=suffix)
        image_data = buffered.getvalue()
        file_data = BytesIO(image_data)

        return file_data

def get_image(url):
    print(f"------get_image:{url}")
    response = None
    try:
        response = requests.get(url, stream=True)
        # response = requests.get(url, verify=False, proxies=None)
    except requests.exceptions.RequestException as e:
        print(f"请求图片 {url} 失败: {e}")

    if not response:
        return None
    elif response.status_code != 200:
        print(f"http请求失败，状态码为{response.status_code}")
        return None
    else:
        image = BytesIO(response.content)
        return image


# 请求url获取响应
def get_response(url):
    print(f"------get_response:{url}")
    response = None
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3",
        "Accept-Language": "en-US,en;q=0.5",
        "Accept-Encoding": "gzip, deflate",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1"
    }
    try:
        response = requests.get(url, headers=headers, verify=False, proxies=None)
    except Exception as e:
        # print(f"http请求失败：{e}")
        print(f"请求页面 {url} 失败: {e}")

    if not response:
        # print(f"http请求失败，返回值为{response}")
        return None
    elif response.status_code != 200:
        print(f"http请求失败，状态码为{response.status_code}")
        return None
    else:
        return response


def set_format(document):
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_reply(text, document):
    # 添加段落
    paragraph = document.add_paragraph()

    # 添加文本
    run = paragraph.add_run(text)

    # 设置文本字体
    font = run.font  # 获取文本的字体对象
    font.name = '宋体'  # 设置字体名称
    font.size = Pt(10)  # 设置字体大小
    # font.bold = True  # 设置字体是否加粗
    font.italic = True  # 设置字体是否倾斜
    # font.underline = True  # 设置字体是否下划线

    # # 设置段落对齐方式
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐