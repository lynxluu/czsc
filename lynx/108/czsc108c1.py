import os

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from threading import Lock, Thread
import io
from datetime import datetime

# 定义爬取的网站地址和页面范围
base_url = 'https://chzhshch.blog/stocks/'
page_range = range(1, 6)

# 创建一个列表，用于保存爬取的文章内容
articles = []

# 创建一个线程锁
lock = Lock()
threads = []

def get_abs_url(rel_url):
    base_url = 'https://chzhshch.blog'
    url = base_url + rel_url
    return url

# 定义一个函数，用于爬取文章内容
def crawl(page_number):
    # page_url = base_url + str(page_number).zfill(3)
    page_url = get_abs_url("/stocks/"+str(page_number).zfill(3))
    try:
        response = requests.get(page_url)
        response.raise_for_status()  # 检查响应状态码是否为 200
    except requests.exceptions.RequestException as e:
        print(f"请求页面 {page_url} 失败: {e}")
        return

    soup = BeautifulSoup(response.content, 'html.parser')
    article = soup.find('article')
    title = article.find('h1').text.strip()

    # 将爬取的文章内容保存到列表中
    with lock:
        articles.append((page_number, title, article.contents))

# 创建线程池，并启动多个线程进行爬取操作
# def main():

for page_number in page_range:
    thread = Thread(target=crawl, args=(page_number,))
    threads.append(thread)
    thread.start()

# 等待所有线程结束，然后对文章内容按照标题进行排序
for thread in threads:
    thread.join()

articles = sorted(articles, key=lambda x: x[0])

# 创建一个Word文档对象, 设置默认字体为微软雅黑
document = Document()

document.styles['Normal'].font.name = '微软雅黑'
document.styles['Normal'].font.size = Pt(12)
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 将爬取的文章内容写入Word文档
for page_number, title, contents in articles:
    document.add_heading(title, level=1)
    # print(type(contents))
    for element in contents[0]:
        # print(element，element)
        if element.name == 'p':
            document.add_paragraph(element.text.strip())
        elif element.name == 'div':
            for img in element.children:
                if img.name == 'img':
                    image_url = img['src']
                    try:
                        # image_response = requests.get(image_url)
                        image_response = requests.get(get_abs_url(image_url),stream=True)
                        image = io.BytesIO(image_response.content)
                        if image_response.raise_for_status() == 200:  # 检查响应状态码是否为 200
                            document.add_picture(image, width=Inches(6))
                    except requests.exceptions.RequestException as e:
                        print(f"请求图片 {image_url} 失败: {e}")


        # https://chzhshch.blog/stocks/002
        # 从这篇来看 h2后面接的平级的若干个div才是真正的回复
        elif element.name in ['h2', 'h3']:
            # print('h2',element)
            document.add_heading('回复', level=2)
            # 获取下一级标签是 <h2> 的内容
            divs = element.find_next_siblings('div')
            # print('divs',divs)
            # end_element = element.find_next_sibling(['h2', 'h3'])
            # if end_element:
            #     divs = divs[:divs.index(end_element)]
            for div in divs:
                # print('div',div)
                # 获取回复者
                replyer = ''
                replyer_tag = div.find('span')
                replyer = replyer_tag.get_text() if replyer_tag else ''
                # print(replyer)

                # 获取回复时间
                time_str = ''
                try:
                    time_str = div.find('br').previous_sibling.strip()
                    # reply_time = datetime.strptime(time_str, '%Y-%m-%d %H:%M:%S')
                    if not time_str:
                        time_str = None
                except Exception as e:
                    print(e)
                # print(time_str)
                document.add_paragraph(replyer+': ' +time_str)

                #获取回复内容
                reply_text = ''
                try:
                    p_tags = div.find_all('p')
                    # 如果p_tags存在，循环拼接p标签内容；如不存在，则取div内容
                    if p_tags:
                        # print(div,p_tags)
                        for p_tag in p_tags:
                            reply_text += p_tag.text.strip() + '\n'
                    else:
                        #替换掉多余的回复人和、回复时间
                        reply_text = div.get_text(strip=True).replace(time_str, '').replace(replyer, '') + '\n'
                except Exception as e:
                    print(e)

                # print(reply_text)
                document.add_paragraph(reply_text)


    # print(page_number, title, contents)

# 保存Word文档
ext = datetime.now().strftime('%y%m%d%_H%M%S')
file_path = os.path.join(f'czsc108c1_{ext}.docx')

document.save(file_path)

# if __name__ == "__main__":
#     main()