from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import etree
from lib108 import *
import concurrent.futures
from queue import Queue





def parse_article1(url, document):
    response = get_response(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # 解析标题
    title = soup.find('h1').text

    # 解析文章内容
    content = ''
    paragraphs = soup.find_all('p')
    images = soup.find_all('img')

    for i in range(len(paragraphs)):
        if i < len(images):
            # 处理图片
            add_image(get_abs_url(images[i]['src']))
            # 将图片下载保存到本地或者直接插入到 Word 文档中
            # 可使用 requests 库下载图片：requests.get(image_url).content
            # 可使用 python-docx 库插入图片：document.add_picture(image_path)

        # 处理文字内容
        content += paragraphs[i].text + '\n'

    # 创建段落并设置标题和内容
    paragraph = document.add_paragraph()
    paragraph.add_run(title).bold = True
    paragraph.add_run('\n' + content)

    # 处理回复内容
    reply_header = soup.find('h2')
    if reply_header:
        # 解析回复内容
        spans = reply_header.find_next_siblings('span')

        for span in spans:
            replies = span.find_all(['p', 'br'])
            for reply in replies:
                # 处理回复的文字内容
                reply_text = reply.text.strip()
                if reply_text:
                    document.add_paragraph(reply_text)

def test1():
    base_url = 'https://chzhshch.blog/stocks/XXX'  # 将XXX替换为具体的页码范围
    document = Document()

    # 遍历页面范围
    for i in range(1, 3):
        url = base_url.replace('XXX', str(i).zfill(3))
        parse_article1(url)

    # 保存 Word 文档
    document.save('result.docx')

# test1()

def add_toc(document):
    p = document.add_paragraph()
    p.add_run().add_break()
    r = p.add_run()

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # fldSimple = etree.Element(etree.QName(nsmap['w'], 'fldSimple'), nsmap=nsmap)
    # fldSimple.set('instr', r'TOC \o "1-3" \h \z \u')

    # 创建一个带有命名空间的 XML 元素
    fldSimple = etree.Element(etree.QName(nsmap['w'], 'fldSimple'),
                              attrib={"w_instr": r'TOC \o "1-3" \h \z \u'}, nsmap=nsmap)


    r.element.append(fldSimple)
    r.add_text("目录")
    r.add_break()

def add_toc2(document):
    p = document.add_paragraph()
    p.add_run().add_break()
    r = p.add_run()

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 创建 XML 元素
    instrText = etree.Element(etree.QName(nsmap['w'], 'instrText'), nsmap=nsmap)
    instrText.text = r'TOC \o "1-3" \h \z \u'

    fldCharBegin = etree.Element(etree.QName(nsmap['w'], 'fldChar'), attrib={"w_fieldCharType": "begin"}, nsmap=nsmap)
    fldCharEnd = etree.Element(etree.QName(nsmap['w'], 'fldChar'), attrib={"w_fieldCharType": "end"}, nsmap=nsmap)

    fldSimple = etree.Element(etree.QName(nsmap['w'], 'fldSimple'), nsmap=nsmap)
    fldSimple.append(fldCharBegin)
    fldSimple.append(instrText)
    fldSimple.append(fldCharEnd)

    # 将 XML 元素添加到 Word 文档中
    r._r.append(fldSimple)
    r.add_text("目录")
    r.add_break()


def validate_toc(document):
    # 在文档中查找所有的目录
    toc_items = document.tables[0].cell(0, 0).paragraphs

    # 遍历目录中的所有章节标题
    for item in toc_items:
        # 如果章节标题的样式不是 Heading 1 或 Heading 2，则跳过
        if item.style.name not in ['Heading 1', 'Heading 2']:
            continue

        # 获取章节标题的文本和页码
        text = item.text.strip()
        page = item.runs[-1].text.strip()

        # 在文档中查找章节标题所在的段落
        paragraphs = document.paragraphs
        found = False
        for p in paragraphs:
            if text in p.text:
                found = True
                break

        # 验证章节标题是否存在于文档中，并且页码是否正确
        if not found:
            print(f"目录验证失败：章节标题“{text}”不存在于文档中")
        elif not page.isdigit() or int(page) != (paragraphs.index(p) + 1):
            print(f"目录验证失败：章节标题“{text}”页码不正确")
        else:
            print(f"目录验证成功：章节标题“{text}”页码为{page}")

def parse_article2(url, document):
    response = get_response(url)
    if not response:
        return

    soup = BeautifulSoup(response.text, 'html.parser')

    # 解析标题
    title = soup.find('h1').text
    document.add_heading(title, level=1)

    # 解析发表时间
    # timestamp = soup.find('blockquote').text
    # document.add_paragraph(timestamp)

    # 解析文章内容
    article = soup.find('article')
    paragraphs = article.find_all('p')
    images = article.find_all('img')

    for i in range(len(paragraphs)):
        if i < len(images):
            # 处理图片
            img_url = get_abs_url(images[i]['src'])
            document.add_paragraph(img_url)
            # add_image(img_url)
        # 处理文字内容
        text = paragraphs[i].text
        document.add_paragraph(text)

    # 解析回复内容
    replies = article.find_all(['h2', 'h3'])
    for reply in replies:
        reply_text = reply.text
        document.add_heading(reply_text, level=2)
        reply_divs = reply.find_next_siblings('div')
        for reply_div in reply_divs:
            reply_time = reply_div.br.previous_sibling.strip()
            reply_author = reply_div.span.text.strip()
            reply_content = reply_div.text.replace(reply_time, '').replace(reply_author, '').strip()
            document.add_paragraph(reply_author + ': ' + reply_time)
            # document.add_paragraph(reply_author)
            if reply_content:
                document.add_paragraph(reply_content)

    # return document


def test2():
    base_url = 'https://chzhshch.blog/stocks/XXX'  # 将XXX替换为具体的页码范围
    document = Document()
    set_format(document)

    toc = document.add_paragraph().add_run().element
    toc.text = "目录"
    toc.style = "Heading1"
    # 遍历页面范围
    for i in range(1, 6):
        url = base_url.replace('XXX', str(i).zfill(3))
        parse_article2(url, document)

    # 添加目录
    # add_toc(document)
    # add_toc2(document)

    # 保存 Word 文档
    document.save('result.docx')

    # validate_toc(document)

test2()




def parse_article3(url):
    document = Document()
    response = get_response(url)
    if not response:
        return

    soup = BeautifulSoup(response.text, 'html.parser')

    # 解析标题
    title = soup.find('h1').text
    document.add_heading(title, level=1)

    # 解析发表时间
    # timestamp = soup.find('blockquote').text
    # document.add_paragraph(timestamp)

    # 解析文章内容
    article = soup.find('article')
    paragraphs = article.find_all('p')
    images = article.find_all('img')

    for i in range(len(paragraphs)):
        if i < len(images):
            # 处理图片
            img_url = get_abs_url(images[i]['src'])
            # document.add_paragraph(img_url)
            document.add_picture(get_image(img_url), width=Inches(5))
        # 处理文字内容
        text = paragraphs[i].text
        document.add_paragraph(text)

    # 解析回复内容
    replies = article.find_all(['h2', 'h3'])
    for reply in replies:
        reply_text = reply.text
        document.add_heading(reply_text, level=2)
        reply_divs = reply.find_next_siblings('div')
        for reply_div in reply_divs:
            reply_time = reply_div.br.previous_sibling.strip()
            reply_author = reply_div.span.text.strip()
            reply_content = reply_div.text.replace(reply_time, '').replace(reply_author, '').strip()
            document.add_paragraph(reply_author + ': ' + reply_time)
            # document.add_paragraph(reply_author)
            if reply_content:
                document.add_paragraph(reply_content)

    return document

def process_page(url, queue):
    # 处理单个页面，将结果放入队列
    result = parse_article3(url)
    queue.put(result)


def test3():
    base_url = 'https://chzhshch.blog/stocks/XXX'  # 将XXX替换为具体的页码范围
    document = Document()

    # 创建队列和线程池
    queue = Queue()
    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
        # 提交任务到线程池
        futures = [executor.submit(process_page, base_url.replace('XXX', str(i).zfill(3)), queue) for i in range(1, 6)]

        # 获取结果并按照原始顺序写入文档
        for future in concurrent.futures.as_completed(futures):
            result = future.result()
            # 将结果写入文档
            write_to_document(result, document)

    # 保存 Word 文档
    document.save('result.docx')

def write_to_document(result, document):
    # 将结果写入文档的相关逻辑
    pass

# test3()



